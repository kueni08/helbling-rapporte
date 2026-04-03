from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def add_code(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shd)
    return p

def add_warning(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('⚠  ' + text)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC5, 0x50, 0x00)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF3CD')
    p._p.get_or_add_pPr().append(shd)
    return p

def add_info(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('ℹ  ' + text)
    run.font.color.rgb = RGBColor(0x0C, 0x54, 0x93)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D1ECF1')
    p._p.get_or_add_pPr().append(shd)
    return p

# ─── TITELSEITE ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MIGRATION HELBLING RAPPORTE')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Hostpoint Hosting')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Schritt-für-Schritt Anleitung')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x44, 0x47, 0x2A)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Erstellt: April 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ─── INHALTSVERZEICHNIS ───────────────────────────────────────────────────────
set_heading(doc, 'Inhaltsverzeichnis', 1)
toc_items = [
    ('1', 'Voraussetzungen & Hosting-Plan'),
    ('2', 'Server-Zugang & Umgebung einrichten'),
    ('3', 'Applikation übertragen'),
    ('4', 'Verzeichnisse & Datenbank einrichten'),
    ('5', 'Umgebungsvariablen konfigurieren (.env)'),
    ('6', 'Node.js Abhängigkeiten installieren'),
    ('7', 'Applikation mit PM2 starten'),
    ('8', 'Nginx Reverse Proxy konfigurieren'),
    ('9', 'SSL / HTTPS einrichten'),
    ('10', 'Domain & DNS bei Hostpoint'),
    ('11', 'Erstanmeldung & Sicherheit'),
    ('12', 'E-Mail (SMTP) konfigurieren'),
    ('13', 'Bestehende Daten migrieren'),
    ('14', 'Fehlerbehebung'),
    ('15', 'Abschluss-Checkliste'),
]
for num, title in toc_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{num}. {title}')
    run.font.size = Pt(11)

doc.add_page_break()

# ─── 1. VORAUSSETZUNGEN ───────────────────────────────────────────────────────
set_heading(doc, '1. Voraussetzungen & Hosting-Plan', 1)
doc.add_paragraph(
    'Die Helbling Rapporte App basiert auf Node.js und benötigt deshalb einen Hosting-Plan '
    'mit SSH-Zugang und Node.js-Unterstützung. Ein einfaches Shared Hosting (PHP) reicht nicht aus.'
)

set_heading(doc, '1.1 Empfohlener Hostpoint-Plan', 2)
doc.add_paragraph('Empfehlung: Hostpoint VPS oder Cloud Server mit folgenden Mindestanforderungen:')

items = [
    'Betriebssystem: Ubuntu 22.04 LTS oder Debian 12',
    'RAM: mindestens 1 GB (empfohlen: 2 GB, wegen Chromium/Puppeteer für PDF-Generierung)',
    'Speicher: mindestens 10 GB SSD',
    'SSH-Zugang (root oder sudo)',
    'Node.js 18 oder höher',
    'Chromium-Browser (für PDF-Generierung)',
    'Öffentliche IPv4-Adresse',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')

add_info(doc, 'Hostpoint bietet VPS-Pläne unter hostpoint.ch/vps an. Alternativ funktioniert auch ein Managed Node.js Hosting, sofern SSH und PM2 verfügbar sind.')

set_heading(doc, '1.2 Technischer Überblick der App', 2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Komponente'
hdr[1].text = 'Details'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
rows_data = [
    ('Laufzeit', 'Node.js 18+'),
    ('Framework', 'Express.js'),
    ('Datenbank', 'SQLite (keine separate DB-Installation nötig!)'),
    ('Port', '3000 (intern, via Nginx nach aussen)'),
    ('PDF-Generierung', 'Puppeteer + Chromium'),
    ('E-Mail', 'SMTP (Outlook / Gmail)'),
    ('Datei-Uploads', 'Lokales Verzeichnis ./uploads/'),
    ('Session-Speicher', 'SQLite (sessions.db)'),
]
for comp, detail in rows_data:
    row = table.add_row().cells
    row[0].text = comp
    row[1].text = detail

doc.add_paragraph()
doc.add_page_break()

# ─── 2. SERVER EINRICHTEN ─────────────────────────────────────────────────────
set_heading(doc, '2. Server-Zugang & Umgebung einrichten', 1)
doc.add_paragraph('Sobald der Hostpoint VPS bereitgestellt ist, verbinden Sie sich via SSH:')
add_code(doc, 'ssh root@IHRE-SERVER-IP')

set_heading(doc, '2.1 System aktualisieren', 2)
add_code(doc, 'apt update && apt upgrade -y')

set_heading(doc, '2.2 Node.js 22 installieren', 2)
doc.add_paragraph('Installation via NodeSource-Repository (empfohlen):')
add_code(doc, 'curl -fsSL https://deb.nodesource.com/setup_22.x | bash -\napt install -y nodejs')
doc.add_paragraph('Version prüfen:')
add_code(doc, 'node --version   # sollte v22.x.x anzeigen\nnpm --version')

set_heading(doc, '2.3 Chromium installieren (für PDF-Generierung)', 2)
add_code(doc, 'apt install -y chromium-browser\nchromium-browser --version')
doc.add_paragraph('Den Pfad notieren (wird später in .env benötigt):')
add_code(doc, 'which chromium-browser\n# Ausgabe: /usr/bin/chromium-browser')

set_heading(doc, '2.4 PM2 installieren (Prozessmanager)', 2)
doc.add_paragraph('PM2 hält die App am Laufen und startet sie bei Serverneustarts automatisch:')
add_code(doc, 'npm install -g pm2')

set_heading(doc, '2.5 Git installieren (optional, für direkten Pull)', 2)
add_code(doc, 'apt install -y git')

set_heading(doc, '2.6 App-Benutzer erstellen (empfohlen)', 2)
doc.add_paragraph('Aus Sicherheitsgründen die App nicht als root betreiben:')
add_code(doc, 'adduser helbling\nusermod -aG sudo helbling\nsu - helbling')

doc.add_page_break()

# ─── 3. APPLIKATION ÜBERTRAGEN ────────────────────────────────────────────────
set_heading(doc, '3. Applikation übertragen', 1)
doc.add_paragraph('Es gibt zwei Möglichkeiten, die Applikation auf den Server zu übertragen:')

set_heading(doc, '3.1 Option A: Via Git (empfohlen)', 2)
doc.add_paragraph('Wenn das Repository auf GitHub/GitLab vorhanden ist:')
add_code(doc, 'cd /home/helbling\ngit clone https://github.com/kueni08/helbling-rapporte.git rapporte\ncd rapporte')

set_heading(doc, '3.2 Option B: Via SFTP / FileZilla', 2)
items = [
    'FileZilla herunterladen und öffnen',
    'Host: IHRE-SERVER-IP, Benutzer: helbling, Passwort: Ihr Passwort, Port: 22',
    'Alle Projektdateien in /home/helbling/rapporte/ hochladen',
    'WICHTIG: Den Ordner node_modules NICHT hochladen (wird lokal installiert)',
    'WICHTIG: Die .env Datei NICHT hochladen (wird auf dem Server erstellt)',
    'WICHTIG: Die db/*.db Dateien nur übertragen, wenn bestehende Daten migriert werden sollen',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

set_heading(doc, '3.3 Ordnerstruktur nach der Übertragung', 2)
add_code(doc, '/home/helbling/rapporte/\n├── server.js\n├── package.json\n├── package-lock.json\n├── routes/\n├── lib/\n├── middleware/\n├── public/\n├── db/          ← wird automatisch erstellt\n├── uploads/     ← wird automatisch erstellt\n└── .env         ← wird manuell erstellt (Schritt 5)')

doc.add_page_break()

# ─── 4. VERZEICHNISSE & DATENBANK ────────────────────────────────────────────
set_heading(doc, '4. Verzeichnisse & Datenbank einrichten', 1)

set_heading(doc, '4.1 Wichtiger Hinweis zur Datenbank', 2)
add_info(doc, 'Die App verwendet SQLite – es muss KEIN separater Datenbankserver (MySQL, PostgreSQL etc.) installiert werden! Die Datenbank ist eine einfache Datei (rapporte.db), die beim ersten Start der App automatisch erstellt und mit allen Tabellen befüllt wird.')

set_heading(doc, '4.2 Verzeichnisse anlegen', 2)
doc.add_paragraph('Folgende Verzeichnisse müssen vorhanden und beschreibbar sein:')
add_code(doc, 'cd /home/helbling/rapporte\nmkdir -p db\nmkdir -p uploads\nmkdir -p uploads/attachments\nmkdir -p uploads/photos\nmkdir -p uploads/lieferscheine-inbox\nmkdir -p uploads/lieferscheine-processed\nmkdir -p logs')

set_heading(doc, '4.3 Berechtigungen setzen', 2)
add_code(doc, 'chown -R helbling:helbling /home/helbling/rapporte\nchmod -R 755 /home/helbling/rapporte\nchmod -R 775 /home/helbling/rapporte/db\nchmod -R 775 /home/helbling/rapporte/uploads')

set_heading(doc, '4.4 Datenbank-Initialisierung', 2)
doc.add_paragraph(
    'Die Datenbank wird beim ersten Start der App (Schritt 7) automatisch erstellt. '
    'Dabei werden folgende Tabellen angelegt:'
)
tables_info = [
    ('users', 'Benutzerkonten (Admin, Planer, Monteure)'),
    ('customers', 'Kundenstammdaten'),
    ('orders', 'Aufträge / Rapporte (Haupttabelle)'),
    ('order_attachments', 'Angehängte Dateien zu Aufträgen'),
    ('order_photos', 'Fotos zu Aufträgen'),
    ('multiselect_options', 'Konfigurierbare Dropdown-Optionen'),
    ('articles', 'Artikelstammdaten'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Tabelle'
hdr[1].text = 'Inhalt'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
for tname, tdesc in tables_info:
    row = table.add_row().cells
    row[0].text = tname
    row[1].text = tdesc

doc.add_paragraph()
add_warning(doc, 'Beim ersten Start werden automatisch Demo-Daten und Standard-Benutzer angelegt. Diese müssen nach der Erstanmeldung sofort geändert werden! (Siehe Schritt 11)')

doc.add_page_break()

# ─── 5. UMGEBUNGSVARIABLEN ────────────────────────────────────────────────────
set_heading(doc, '5. Umgebungsvariablen konfigurieren (.env)', 1)
doc.add_paragraph(
    'Die App wird über eine .env Datei konfiguriert. Diese Datei enthält alle sensiblen '
    'Einstellungen und darf NICHT ins Git-Repository eingecheckt werden.'
)

set_heading(doc, '5.1 .env Datei erstellen', 2)
add_code(doc, 'cd /home/helbling/rapporte\nnano .env')

doc.add_paragraph('Folgenden Inhalt einfügen und anpassen:')
add_code(doc, '# ── Server ──────────────────────────────────\nPORT=3000\nNODE_ENV=production\nSESSION_SECRET=HIER-EINEN-LANGEN-ZUFAELLIGEN-STRING-EINFUEGEN\n\n# ── Chromium (PDF-Generierung) ──────────────\nCHROMIUM_PATH=/usr/bin/chromium-browser\n\n# ── Datei-Uploads ───────────────────────────\nUPLOADS_DIR=./uploads\n\n# ── E-Mail (SMTP) ───────────────────────────\nSMTP_HOST=smtp.office365.com\nSMTP_PORT=587\nSMTP_SECURE=false\nSMTP_USER=rapport@ihre-firma.ch\nSMTP_PASS=IHR-E-MAIL-PASSWORT\n\n# ── Anthropic Claude API (optional) ─────────\n# ANTHROPIC_API_KEY=sk-ant-...\n\n# ── Google Drive (optional) ─────────────────\n# GOOGLE_DRIVE_FOLDER_ID=\n# GOOGLE_OAUTH_CLIENT_ID=\n# GOOGLE_OAUTH_CLIENT_SECRET=\n# GOOGLE_OAUTH_REFRESH_TOKEN=')

doc.add_paragraph('Datei speichern: Ctrl+O, Enter, Ctrl+X')

set_heading(doc, '5.2 SESSION_SECRET generieren', 2)
doc.add_paragraph('Einen sicheren, zufälligen String generieren:')
add_code(doc, 'node -e "console.log(require(\'crypto\').randomBytes(64).toString(\'hex\'))"')
doc.add_paragraph('Den ausgegebenen String in .env als SESSION_SECRET eintragen.')

set_heading(doc, '5.3 Übersicht aller Umgebungsvariablen', 2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(['Variable', 'Beispielwert', 'Beschreibung', 'Pflicht']):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(9)
env_rows = [
    ('PORT', '3000', 'Interner App-Port', 'Nein'),
    ('NODE_ENV', 'production', 'Betriebsmodus', 'Nein'),
    ('SESSION_SECRET', '(langer Zufallsstring)', 'Session-Verschlüsselung', 'JA'),
    ('CHROMIUM_PATH', '/usr/bin/chromium-browser', 'Pfad zu Chromium', 'JA'),
    ('UPLOADS_DIR', './uploads', 'Upload-Verzeichnis', 'Nein'),
    ('SMTP_HOST', 'smtp.office365.com', 'Mail-Server', 'Für E-Mail'),
    ('SMTP_PORT', '587', 'Mail-Port', 'Für E-Mail'),
    ('SMTP_SECURE', 'false', 'TLS direkt (465) oder STARTTLS (587)', 'Für E-Mail'),
    ('SMTP_USER', 'rapport@firma.ch', 'Absender-E-Mail', 'Für E-Mail'),
    ('SMTP_PASS', '(Passwort)', 'SMTP-Passwort', 'Für E-Mail'),
    ('ANTHROPIC_API_KEY', 'sk-ant-...', 'Claude API (Lieferschein-Import)', 'Optional'),
    ('GOOGLE_DRIVE_FOLDER_ID', '(Ordner-ID)', 'Google Drive Zielordner', 'Optional'),
]
for v, ex, desc, req in env_rows:
    row = table.add_row().cells
    row[0].text = v
    row[1].text = ex
    row[2].text = desc
    row[3].text = req
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
doc.add_page_break()

# ─── 6. NPM INSTALL ───────────────────────────────────────────────────────────
set_heading(doc, '6. Node.js Abhängigkeiten installieren', 1)
doc.add_paragraph('Im App-Verzeichnis alle Pakete installieren:')
add_code(doc, 'cd /home/helbling/rapporte\nnpm install --omit=dev')
doc.add_paragraph('Dies kann einige Minuten dauern, da puppeteer-core und andere Pakete heruntergeladen werden.')

set_heading(doc, '6.1 Installation prüfen', 2)
add_code(doc, 'ls node_modules | wc -l   # sollte > 100 Pakete anzeigen')

set_heading(doc, '6.2 Chromium-Pfad verifizieren', 2)
add_code(doc, 'ls -la $(which chromium-browser)\n# oder:\nls -la /usr/bin/chromium-browser')
add_warning(doc, 'Falls chromium-browser nicht unter /usr/bin/chromium-browser liegt, den korrekten Pfad in der .env Datei unter CHROMIUM_PATH eintragen.')

doc.add_page_break()

# ─── 7. PM2 STARTEN ───────────────────────────────────────────────────────────
set_heading(doc, '7. Applikation mit PM2 starten', 1)

set_heading(doc, '7.1 PM2 Konfigurationsdatei erstellen', 2)
doc.add_paragraph('Eine ecosystem.config.js Datei im App-Verzeichnis erstellen:')
add_code(doc, 'nano /home/helbling/rapporte/ecosystem.config.js')
doc.add_paragraph('Inhalt:')
add_code(doc, "module.exports = {\n  apps: [{\n    name: 'helbling-rapporte',\n    script: 'server.js',\n    cwd: '/home/helbling/rapporte',\n    instances: 1,\n    autorestart: true,\n    watch: false,\n    max_memory_restart: '512M',\n    env: {\n      NODE_ENV: 'production',\n      PORT: 3000\n    },\n    error_file: './logs/pm2-error.log',\n    out_file: './logs/pm2-out.log',\n    log_date_format: 'YYYY-MM-DD HH:mm:ss'\n  }]\n};")

set_heading(doc, '7.2 App starten', 2)
add_code(doc, 'cd /home/helbling/rapporte\npm2 start ecosystem.config.js')

set_heading(doc, '7.3 Autostart bei Serverstart einrichten', 2)
add_code(doc, 'pm2 startup\n# Den ausgegebenen Befehl kopieren und ausführen (beginnt mit: sudo env PATH=...)\npm2 save')

set_heading(doc, '7.4 Status und Logs prüfen', 2)
add_code(doc, 'pm2 status                          # App-Status anzeigen\npm2 logs helbling-rapporte         # Live-Logs\npm2 logs helbling-rapporte --lines 50  # Letzte 50 Zeilen')

doc.add_paragraph('Die App sollte nun auf Port 3000 laufen. Test:')
add_code(doc, 'curl http://localhost:3000\n# sollte HTML zurückgeben')

set_heading(doc, '7.5 Nützliche PM2-Befehle', 2)
add_code(doc, 'pm2 restart helbling-rapporte   # App neu starten\npm2 stop helbling-rapporte      # App stoppen\npm2 delete helbling-rapporte    # App aus PM2 entfernen\npm2 monit                       # Ressourcen-Monitor')

doc.add_page_break()

# ─── 8. NGINX ─────────────────────────────────────────────────────────────────
set_heading(doc, '8. Nginx Reverse Proxy konfigurieren', 1)
doc.add_paragraph(
    'Nginx leitet Anfragen von Port 80/443 (HTTP/HTTPS) an die App auf Port 3000 weiter. '
    'So ist die App unter einer normalen URL erreichbar.'
)

set_heading(doc, '8.1 Nginx installieren', 2)
add_code(doc, 'apt install -y nginx\nsystemctl enable nginx\nsystemctl start nginx')

set_heading(doc, '8.2 Nginx-Konfiguration erstellen', 2)
add_code(doc, 'nano /etc/nginx/sites-available/helbling-rapporte')
doc.add_paragraph('Inhalt (IHRE-DOMAIN durch echte Domain ersetzen):')
add_code(doc, 'server {\n    listen 80;\n    server_name IHRE-DOMAIN.ch www.IHRE-DOMAIN.ch;\n\n    # Maximale Upload-Grösse (für Foto-Uploads)\n    client_max_body_size 50M;\n\n    location / {\n        proxy_pass http://127.0.0.1:3000;\n        proxy_http_version 1.1;\n        proxy_set_header Upgrade $http_upgrade;\n        proxy_set_header Connection \'upgrade\';\n        proxy_set_header Host $host;\n        proxy_set_header X-Real-IP $remote_addr;\n        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;\n        proxy_set_header X-Forwarded-Proto $scheme;\n        proxy_cache_bypass $http_upgrade;\n        proxy_read_timeout 300s;\n        proxy_connect_timeout 75s;\n    }\n\n    # Logs\n    access_log /var/log/nginx/rapporte-access.log;\n    error_log  /var/log/nginx/rapporte-error.log;\n}')

set_heading(doc, '8.3 Konfiguration aktivieren', 2)
add_code(doc, 'ln -s /etc/nginx/sites-available/helbling-rapporte /etc/nginx/sites-enabled/\nnginx -t              # Konfiguration testen\nsystemctl reload nginx')

doc.add_page_break()

# ─── 9. SSL ───────────────────────────────────────────────────────────────────
set_heading(doc, '9. SSL / HTTPS einrichten', 1)
doc.add_paragraph(
    'SSL-Zertifikat mit Let\'s Encrypt (kostenlos) einrichten. '
    'Voraussetzung: Domain muss bereits auf den Server zeigen (Schritt 10).'
)

set_heading(doc, '9.1 Certbot installieren', 2)
add_code(doc, 'apt install -y certbot python3-certbot-nginx')

set_heading(doc, '9.2 Zertifikat erstellen', 2)
add_code(doc, 'certbot --nginx -d IHRE-DOMAIN.ch -d www.IHRE-DOMAIN.ch')
doc.add_paragraph('Certbot aktualisiert die Nginx-Konfiguration automatisch für HTTPS.')

set_heading(doc, '9.3 Automatische Erneuerung prüfen', 2)
add_code(doc, 'certbot renew --dry-run')
add_info(doc, 'Let\'s Encrypt Zertifikate sind 90 Tage gültig. Certbot erneuert sie automatisch via systemd-Timer.')

set_heading(doc, '9.4 Alternative: Hostpoint SSL', 2)
doc.add_paragraph(
    'Hostpoint bietet in der Verwaltungsoberfläche auch eigene SSL-Zertifikate an. '
    'Diese können alternativ zu Let\'s Encrypt verwendet werden. '
    'In diesem Fall das Zertifikat in Nginx manuell konfigurieren.'
)

doc.add_page_break()

# ─── 10. DOMAIN & DNS ─────────────────────────────────────────────────────────
set_heading(doc, '10. Domain & DNS bei Hostpoint', 1)
doc.add_paragraph('Die Domain muss auf die IP-Adresse des Servers zeigen.')

set_heading(doc, '10.1 DNS-Einstellungen', 2)
doc.add_paragraph('Im Hostpoint-Kundencenter (my.hostpoint.ch) folgende DNS-Einträge setzen:')
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(['Typ', 'Name', 'Wert', 'TTL']):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.font.bold = True
dns_rows = [
    ('A', '@', 'IHRE-SERVER-IP', '3600'),
    ('A', 'www', 'IHRE-SERVER-IP', '3600'),
    ('AAAA', '@', 'IHRE-IPv6 (falls vorhanden)', '3600'),
]
for t, n, v, ttl in dns_rows:
    row = table.add_row().cells
    for i, val in enumerate([t, n, v, ttl]):
        row[i].text = val

doc.add_paragraph()
add_info(doc, 'DNS-Änderungen können bis zu 24-48 Stunden dauern bis sie weltweit wirksam sind. In der Schweiz meist innerhalb von 1-2 Stunden.')

set_heading(doc, '10.2 DNS-Propagation prüfen', 2)
add_code(doc, 'nslookup IHRE-DOMAIN.ch\n# oder\ndig IHRE-DOMAIN.ch')

doc.add_page_break()

# ─── 11. ERSTANMELDUNG & SICHERHEIT ──────────────────────────────────────────
set_heading(doc, '11. Erstanmeldung & Sicherheit', 1)
add_warning(doc, 'WICHTIG: Nach dem ersten Start müssen alle Standard-Passwörter sofort geändert werden!')

set_heading(doc, '11.1 Standard-Zugangsdaten', 2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(['Benutzername', 'Standard-Passwort', 'Rolle']):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.font.bold = True
users_data = [
    ('admin', 'admin123', 'Administrator'),
    ('thomas.mueller', 'monteur123', 'Monteur'),
    ('stefan.weber', 'monteur123', 'Monteur'),
    ('markus.huber', 'monteur123', 'Monteur'),
    ('daniel.schmid', 'monteur123', 'Monteur'),
]
for u, p, r in users_data:
    row = table.add_row().cells
    for i, val in enumerate([u, p, r]):
        row[i].text = val

doc.add_paragraph()
add_warning(doc, 'Diese Demo-Benutzer und Passwörter MÜSSEN nach der Erstanmeldung geändert werden. Nicht benötigte Demo-Benutzer löschen oder deaktivieren.')

set_heading(doc, '11.2 Admin-Passwort ändern', 2)
items_sec = [
    'Browser öffnen und zur App-URL navigieren',
    'Mit admin / admin123 einloggen',
    'Oben rechts auf "Admin" klicken → Passwort ändern',
    'Sicheres Passwort mit mind. 12 Zeichen setzen',
]
for item in items_sec:
    doc.add_paragraph(item, style='List Number')

set_heading(doc, '11.3 Monteur-Benutzer verwalten', 2)
doc.add_paragraph('Im Admin-Bereich unter "Benutzer":')
items_usr = [
    'Demo-Monteure löschen oder umbenennen',
    'Echte Mitarbeiter mit korrekten Namen anlegen',
    'Sichere Passwörter vergeben',
    'E-Mail-Adressen der Benutzer hinterlegen',
]
for item in items_usr:
    doc.add_paragraph(item, style='List Bullet')

set_heading(doc, '11.4 Firewall einrichten', 2)
doc.add_paragraph('Nur notwendige Ports öffnen:')
add_code(doc, 'ufw allow 22/tcp    # SSH\nufw allow 80/tcp    # HTTP\nufw allow 443/tcp   # HTTPS\nufw enable\nufw status')
add_info(doc, 'Port 3000 (Node.js App) bleibt geschlossen – der Zugriff erfolgt nur intern via Nginx.')

doc.add_page_break()

# ─── 12. E-MAIL ───────────────────────────────────────────────────────────────
set_heading(doc, '12. E-Mail (SMTP) konfigurieren', 1)

set_heading(doc, '12.1 Outlook / Office 365', 2)
add_code(doc, 'SMTP_HOST=smtp.office365.com\nSMTP_PORT=587\nSMTP_SECURE=false\nSMTP_USER=rapport@ihre-firma.ch\nSMTP_PASS=IHR-PASSWORT')
add_info(doc, 'Bei Office 365 mit aktivierter Multi-Faktor-Authentifizierung (MFA) muss ein App-Passwort erstellt werden: Microsoft 365 Admin Center → Benutzer → App-Passwörter.')

set_heading(doc, '12.2 Gmail', 2)
add_code(doc, 'SMTP_HOST=smtp.gmail.com\nSMTP_PORT=587\nSMTP_SECURE=false\nSMTP_USER=rapport@gmail.com\nSMTP_PASS=IHR-APP-PASSWORT')
add_info(doc, 'Bei Gmail muss ein App-Passwort erstellt werden: Google-Konto → Sicherheit → 2-Schritt-Verifizierung → App-Passwörter.')

set_heading(doc, '12.3 E-Mail-Verbindung testen', 2)
add_code(doc, 'cd /home/helbling/rapporte\nnode -e "\nconst nodemailer = require(\'nodemailer\');\nconst t = nodemailer.createTransport({\n  host: process.env.SMTP_HOST,\n  port: process.env.SMTP_PORT,\n  auth: { user: process.env.SMTP_USER, pass: process.env.SMTP_PASS }\n});\nt.verify().then(() => console.log(\'SMTP OK\')).catch(console.error);\n" --env-file .env')

doc.add_page_break()

# ─── 13. DATENMIGRATION ───────────────────────────────────────────────────────
set_heading(doc, '13. Bestehende Daten migrieren', 1)
doc.add_paragraph(
    'Falls bereits eine bestehende Instanz der App läuft (z.B. auf Fly.io), '
    'können die Daten wie folgt übertragen werden.'
)

set_heading(doc, '13.1 Datenbank exportieren (von altem Server)', 2)
add_code(doc, '# Auf dem alten Server:\ncp /app/db/rapporte.db /tmp/rapporte_backup.db\n\n# Lokal herunterladen:\nscp user@ALTER-SERVER:/tmp/rapporte_backup.db ./rapporte_backup.db')

set_heading(doc, '13.2 Datenbank auf neuem Server einspielen', 2)
add_warning(doc, 'Die App muss gestoppt sein, bevor die Datenbank ersetzt wird!')
add_code(doc, '# App stoppen:\npm2 stop helbling-rapporte\n\n# Backup hochladen:\nscp rapporte_backup.db helbling@NEUER-SERVER:/home/helbling/rapporte/db/rapporte.db\n\n# Berechtigungen setzen:\nchown helbling:helbling /home/helbling/rapporte/db/rapporte.db\nchmod 664 /home/helbling/rapporte/db/rapporte.db\n\n# App wieder starten:\npm2 start helbling-rapporte')

set_heading(doc, '13.3 Uploads / Fotos migrieren', 2)
add_code(doc, '# Auf altem Server: ZIP erstellen\ntar -czf /tmp/uploads_backup.tar.gz /app/uploads/\n\n# Lokal herunterladen:\nscp user@ALTER-SERVER:/tmp/uploads_backup.tar.gz ./\n\n# Auf neuem Server hochladen und entpacken:\nscp uploads_backup.tar.gz helbling@NEUER-SERVER:/home/helbling/rapporte/\nssh helbling@NEUER-SERVER\ncd /home/helbling/rapporte\ntar -xzf uploads_backup.tar.gz --strip-components=2 -C uploads/')

set_heading(doc, '13.4 Datenbank-Integrität prüfen', 2)
add_code(doc, 'cd /home/helbling/rapporte\nnode -e "\nconst db = require(\'./lib/database\');\nconst count = db.prepare(\'SELECT COUNT(*) as n FROM orders\').get();\nconsole.log(\'Aufträge in DB:\', count.n);\n"')

doc.add_page_break()

# ─── 14. FEHLERBEHEBUNG ───────────────────────────────────────────────────────
set_heading(doc, '14. Fehlerbehebung', 1)

set_heading(doc, '14.1 App startet nicht', 2)
add_code(doc, 'pm2 logs helbling-rapporte --lines 50')
items_err = [
    'Port bereits belegt: PORT in .env ändern oder anderen Prozess stoppen (lsof -i :3000)',
    '.env Datei fehlt oder falsch: Prüfen ob .env im App-Verzeichnis liegt',
    'node_modules fehlen: npm install nochmals ausführen',
    'Fehlende Berechtigung: chown -R helbling:helbling /home/helbling/rapporte',
]
for item in items_err:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

set_heading(doc, '14.2 PDF-Generierung schlägt fehl', 2)
items_pdf = [
    'Chromium-Pfad falsch: which chromium-browser ausführen, Pfad in .env anpassen',
    'Chromium nicht installiert: apt install -y chromium-browser',
    'Fehlende Bibliotheken: apt install -y libnss3 libatk-bridge2.0-0 libdrm2 libxcomposite1 libxdamage1',
]
for item in items_pdf:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

set_heading(doc, '14.3 E-Mail-Versand schlägt fehl', 2)
items_mail = [
    'SMTP-Zugangsdaten prüfen (User/Passwort korrekt?)',
    'App-Passwort verwenden statt Account-Passwort (bei MFA)',
    'Port 587 im Hostpoint-Firewall freigegeben?',
    'SMTP_SECURE=false bei Port 587 (STARTTLS)',
]
for item in items_mail:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

set_heading(doc, '14.4 Nginx gibt 502 Bad Gateway', 2)
items_nginx = [
    'App läuft nicht: pm2 status prüfen',
    'Falscher Port in Nginx-Config (muss 3000 sein)',
    'App-Logs prüfen: pm2 logs helbling-rapporte',
]
for item in items_nginx:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

set_heading(doc, '14.5 Datei-Uploads schlagen fehl', 2)
items_up = [
    'Berechtigungen: chmod -R 775 /home/helbling/rapporte/uploads',
    'Speicherplatz: df -h prüfen',
    'Nginx client_max_body_size erhöhen (aktuell: 50M)',
]
for item in items_up:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

doc.add_page_break()

# ─── 15. CHECKLISTE ───────────────────────────────────────────────────────────
set_heading(doc, '15. Abschluss-Checkliste', 1)
doc.add_paragraph('Alle Punkte vor dem Go-Live abhaken:')

checklist_sections = {
    'Server & Umgebung': [
        'Hostpoint VPS/Cloud Server bestellt und SSH-Zugang erhalten',
        'Ubuntu/Debian System aktualisiert (apt update && apt upgrade)',
        'Node.js 22 installiert und Version geprüft',
        'Chromium-Browser installiert',
        'PM2 global installiert',
        'App-Benutzer "helbling" erstellt',
        'Firewall konfiguriert (Ports 22, 80, 443)',
    ],
    'Applikation': [
        'Projektdateien auf Server übertragen (Git oder SFTP)',
        'Verzeichnisse db/, uploads/ und Unterordner erstellt',
        'Dateiberechtigungen korrekt gesetzt',
        '.env Datei erstellt und alle Variablen konfiguriert',
        'SESSION_SECRET mit sicherem Zufallsstring gesetzt',
        'CHROMIUM_PATH auf korrekten Pfad gesetzt',
        'npm install --omit=dev erfolgreich ausgeführt',
    ],
    'Datenbank': [
        'App einmal gestartet (Datenbank wird automatisch initialisiert)',
        'Tabellen in rapporte.db vorhanden geprüft',
        'Bestehende Daten migriert (falls vorhanden)',
        'Datenbankdatei liegt in db/ und ist beschreibbar',
    ],
    'Webserver & Domain': [
        'Nginx installiert und konfiguriert',
        'Nginx-Konfiguration getestet (nginx -t)',
        'DNS-Einträge bei Hostpoint gesetzt (A-Record)',
        'DNS-Propagation abgewartet und geprüft',
        'SSL-Zertifikat mit Certbot erstellt',
        'HTTPS funktioniert (https://IHRE-DOMAIN.ch)',
        'HTTP leitet automatisch auf HTTPS um',
    ],
    'App & Sicherheit': [
        'PM2 Autostart konfiguriert (pm2 startup && pm2 save)',
        'App unter https://IHRE-DOMAIN.ch erreichbar',
        'Login mit admin / admin123 funktioniert',
        'Admin-Passwort geändert',
        'Demo-Monteure gelöscht / echte Benutzer angelegt',
        'Monteur-Passwörter geändert',
        'E-Mail-Versand getestet (Test-Rapport senden)',
        'PDF-Generierung getestet',
        'Foto-Upload getestet',
    ],
    'Backup & Monitoring': [
        'Regelmässiges Datenbank-Backup eingerichtet (Cron-Job)',
        'PM2 Logs werden geschrieben (logs/ Verzeichnis)',
        'Server-Monitoring eingerichtet (optional)',
    ],
}

for section_title, items in checklist_sections.items():
    set_heading(doc, section_title, 2)
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('☐  ' + item)

doc.add_page_break()

# ─── BACKUP CRON ──────────────────────────────────────────────────────────────
set_heading(doc, 'Anhang: Automatisches Datenbank-Backup', 1)
doc.add_paragraph('Tägliches Backup der SQLite-Datenbank einrichten:')
add_code(doc, 'crontab -e')
doc.add_paragraph('Folgende Zeile hinzufügen (täglich um 03:00 Uhr):')
add_code(doc, '0 3 * * * cp /home/helbling/rapporte/db/rapporte.db /home/helbling/backups/rapporte_$(date +\\%Y\\%m\\%d).db && find /home/helbling/backups/ -name "rapporte_*.db" -mtime +30 -delete')
doc.add_paragraph('Backup-Verzeichnis erstellen:')
add_code(doc, 'mkdir -p /home/helbling/backups')
add_info(doc, 'Dieses Backup speichert die Datenbank täglich und löscht Backups die älter als 30 Tage sind automatisch.')

# ─── SPEICHERN ────────────────────────────────────────────────────────────────
output_path = '/home/user/helbling-rapporte/Migration_Hostpoint.docx'
doc.save(output_path)
print(f'Dokument gespeichert: {output_path}')
